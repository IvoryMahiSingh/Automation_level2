import docx

def generate_sdd_document():
    # Create a new Word document
    document = docx.Document()

    # Add the title
    title = document.add_heading('Adhoc Check Solution Design Document', 0)

    # Add the use case and RPA tool
    document.add_paragraph('Use Case: Adhoc Check')
    document.add_paragraph('RPA Tool: Blueprism')

    # Add the steps
    document.add_heading('Steps', 1)
    steps = [
        {'step': '1', 'activity': 'Receive Email and Extract Incident Number', 'description': 'Read email from Automation Bot’s mail-box and extract incident number from email body'},
        {'step': '2', 'activity': 'Initiate Automation and Authenticate Agent', 'description': 'Initiate automation by calling API setup on agent’s machine and authenticate agent’s credentials'},
        {'step': '3', 'activity': 'Validate Request and Get SAP SID and Client Details', 'description': 'Validate request by verifying incident number and CR number against RM team’s email and get SAP SID and client details'},
        {'step': '4', 'activity': 'Open SAP System and Perform Client Opening', 'description': 'Open SAP system using SAP SID and client details and perform client opening steps'},
        {'step': '5', 'activity': 'Wait for Client Closing Signal', 'description': 'Wait for client closing signal from agent or time limit exceeds'},
        {'step': '6', 'activity': 'Close SAP Client and Capture Events', 'description': 'Close SAP client and capture events using STAD report'}
    ]

    for step in steps:
        document.add_paragraph(f'Step {step["step"]}: {step["activity"]}')
        document.add_paragraph(step['description'])

    # Add the business/system exceptions
    document.add_heading('Business/System Exceptions', 1)
    exceptions = [
        {'description': 'Receive Email', 'expected_outcome': 'Email should be received successfully', 'exception_message': 'Failed to receive email. Ensure email account is configured correctly.', 'exception_category': 'System'},
        {'description': 'Authenticate Agent', 'expected_outcome': 'Agent should be authenticated successfully', 'exception_message': 'Authentication failed. Ensure agent credentials are correct.', 'exception_category': 'Business'},
        {'description': 'Validate Request', 'expected_outcome': 'Request should be validated successfully', 'exception_message': 'Request validation failed. Ensure incident number and CR number are correct.', 'exception_category': 'Business'},
        {'description': 'Get SAP SID and Client', 'expected_outcome': 'SAP SID and Client should be extracted successfully', 'exception_message': 'Failed to extract SAP SID and Client. Ensure inputs are correct.', 'exception_category': 'Business'},
        {'description': 'Open SAP System', 'expected_outcome': 'SAP system should be opened successfully', 'exception_message': 'Failed to open SAP system. Ensure SAP SID and Client are correct.', 'exception_category': 'System'},
        {'description': 'Perform Client Opening', 'expected_outcome': 'Client opening should be performed successfully', 'exception_message': 'Failed to perform client opening. Ensure SAP system is accessible.', 'exception_category': 'System'},
        {'description': 'Wait for Client Closing Signal', 'expected_outcome': 'Client closing signal should be received successfully', 'exception_message': 'Failed to receive client closing signal. Ensure agent input is correct.', 'exception_category': 'Business'},
        {'description': 'Close SAP Client', 'expected_outcome': 'SAP client should be closed successfully', 'exception_message': 'Failed to close SAP client. Ensure SAP system is accessible.', 'exception_category': 'System'},
        {'description': 'Capture Events', 'expected_outcome': 'Events should be captured successfully', 'exception_message': 'Failed to capture events. Ensure STAD report is extracted correctly.', 'exception_category': 'System'}
    ]

    table = document.add_table(rows=len(exceptions) + 1, cols=4, style='Table Grid')
    table.cell(0, 0).text = 'Description'
    table.cell(0, 1).text = 'Expected Outcome'
    table.cell(0, 2).text = 'Exception Message'
    table.cell(0, 3).text = 'Exception Category'

    for i, exception in enumerate(exceptions):
        table.cell(i+1, 0).text = exception['description']
        table.cell(i+1, 1).text = exception['expected_outcome']
        table.cell(i+1, 2).text = exception['exception_message']
        table.cell(i+1, 3).text = exception['exception_category']

    # Save the document
    document.save('Adhoc Check Solution Design Document2.docx')

generate_sdd()
